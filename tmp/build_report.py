from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.shape import WD_INLINE_SHAPE


ROOT = Path(r"E:\project\世界稀疏建模\世界稀疏建模01")
SOURCE = ROOT / "docs" / "semantic_world_model_proposal.md"
OUTPUT = ROOT / "outputs" / "semantic_world_model_proposal.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
MUTED = "666666"
TABLE_WIDTH_DXA = 9360


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths_dxa):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string("333333")
    code.paragraph_format.left_indent = Inches(0.22)
    code.paragraph_format.right_indent = Inches(0.22)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(5)
    code.paragraph_format.line_spacing = 1.0

    quote = doc.styles.add_style("Lead Callout", WD_STYLE_TYPE.PARAGRAPH)
    quote.font.name = "Microsoft YaHei"
    quote._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    quote.font.size = Pt(11)
    quote.font.color.rgb = RGBColor.from_string(INK)
    quote.font.bold = True
    quote.paragraph_format.left_indent = Inches(0.15)
    quote.paragraph_format.right_indent = Inches(0.15)
    quote.paragraph_format.space_before = Pt(5)
    quote.paragraph_format.space_after = Pt(8)
    quote.paragraph_format.line_spacing = 1.2

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("世界稀疏建模 | 概念设计基线")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("2026-08-27  |  语义世界模型方案")
    set_run_font(run, size=8.5, color=MUTED)


def add_markdown_runs(paragraph, text, size=10.5, color="000000", bold_default=False, italic=False):
    chunks = re.split(r"(\*\*.*?\*\*)", text)
    for chunk in chunks:
        if not chunk:
            continue
        bold = bold_default
        if chunk.startswith("**") and chunk.endswith("**"):
            chunk = chunk[2:-2]
            bold = True
        run = paragraph.add_run(chunk)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_paragraph(doc, text, style=None, align=None, lead=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if lead:
        p.style = "Lead Callout"
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), CALLOUT)
        p_pr.append(shd)
    add_markdown_runs(p, text, size=11 if lead else 10.5, color=INK if lead else "000000")
    return p


def add_architecture_diagram(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(5)
    add_markdown_runs(title, "四层地图与任务执行闭环", size=13, color=INK, bold_default=True)
    stages = [
        ("传感输入", "RGB-D 双目相机 + IMU", "DDEBF7"),
        ("长期定位", "视觉惯性定位、稀疏地标、关键帧、回环位姿图", "E8EEF5"),
        ("长期语义", "对象、地点、粗锚点、截图证据、语义索引", "EAF4EA"),
        ("长期通行", "房间、门口、转折点、楼梯、电梯、连通边", "F3E8F8"),
        ("短期执行", "局部稠密点云、局部代价地图、避障与精细对准", "FFF1D6"),
        ("任务闭环", "检索候选 -> 选可观察目标位姿 -> 全局拓扑路由 -> 局部安全执行", "DDEBF7"),
    ]
    table = doc.add_table(rows=len(stages) * 2 - 1, cols=2)
    set_table_geometry(table, [2150, 7210])
    for idx, (label, detail, fill) in enumerate(stages):
        row = table.rows[idx * 2]
        for cell in row.cells:
            set_cell_shading(cell, fill)
        for cell, text, bold in ((row.cells[0], label, True), (row.cells[1], detail, False)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_markdown_runs(p, text, size=9.5, color=INK, bold_default=bold)
        if idx < len(stages) - 1:
            arrow = table.rows[idx * 2 + 1]
            merged = arrow.cells[0].merge(arrow.cells[1])
            p = merged.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            add_markdown_runs(p, "↓", size=12, color="456B86", bold_default=True)
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(7)
    add_markdown_runs(caption, "图 1. 长期稀疏世界模型与短期稠密执行缓存分离", size=9, color=MUTED, italic=True)


def is_table_separator(line):
    return bool(re.fullmatch(r"\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?", line))


def split_table(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_table(doc, rows):
    data = [split_table(row) for row in rows if not is_table_separator(row)]
    if not data:
        return
    cols = max(len(row) for row in data)
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Table Grid"
    widths = [TABLE_WIDTH_DXA // cols] * cols
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    for row_idx, row in enumerate(data):
        cells = table.add_row().cells
        for col_idx in range(cols):
            cell = cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[col_idx] if col_idx < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_markdown_runs(p, text, size=8.7, color="000000", bold_default=(row_idx == 0))
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
        set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FB")
    p = cell.paragraphs[0]
    p.style = "Code Block"
    for index, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.3, color="2B2B2B")
        if index < len(lines) - 1:
            run.add_break()
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(76)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("技术方案")
    set_run_font(r, size=12, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("单 RGB-D 双目相机 + IMU\n语义世界稀疏三维建模与导航")
    set_run_font(r, size=25, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("概念设计、原型实施与验证基线")
    set_run_font(r, size=14, color=MUTED)

    table = doc.add_table(rows=3, cols=2)
    set_table_geometry(table, [2700, 6660])
    rows = [("文档版本", "0.2"), ("适用范围", "室内语义建图、跨房间/楼层导航、对象检索"), ("日期", "2026-08-27")]
    for i, (left, right) in enumerate(rows):
        for j, text in enumerate((left, right)):
            cell = table.cell(i, j)
            set_cell_shading(cell, LIGHT_GRAY if j == 0 else "FFFFFF")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_markdown_runs(p, text, size=10, color=INK, bold_default=(j == 0))

    doc.add_paragraph()
    note = doc.add_paragraph(style="Lead Callout")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_markdown_runs(note, "核心判断：地图内全局坐标可由视觉、深度、IMU 与回环稳定维护；绝对建筑坐标仍需要外部基准。", size=11, color=INK, bold_default=True)
    doc.add_page_break()


def add_toc(doc, headings):
    h = doc.add_paragraph(style="Heading 1")
    add_markdown_runs(h, "阅读导航", size=16, color=BLUE, bold_default=True)
    for level, text in headings:
        if level != 2:
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        add_markdown_runs(p, text, size=10.5, color=INK)
    doc.add_page_break()


def build_document():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style_document(doc)
    add_title_page(doc)

    headings = []
    for line in lines:
        if line.startswith("## "):
            headings.append((2, line[3:]))
    add_toc(doc, headings)

    i = 0
    inserted_figure = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_markdown_runs(p, line[3:], size=16, color=BLUE, bold_default=True)
            if line[3:].startswith("2. 总体架构") and not inserted_figure:
                add_architecture_diagram(doc)
                inserted_figure = True
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_markdown_runs(p, line[4:], size=13, color=BLUE, bold_default=True)
            i += 1
            continue
        if line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_markdown_runs(p, line[5:], size=12, color=DARK_BLUE, bold_default=True)
            i += 1
            continue
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue
        if re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            add_markdown_runs(p, line[2:], size=10.5)
            i += 1
            continue
        match = re.match(r"^(\d+)\. (.*)", line)
        if match:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            add_markdown_runs(p, match.group(2), size=10.5)
            i += 1
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            candidate = lines[i]
            if not candidate.strip() or candidate.startswith(("#", "```", "|")) or re.match(r"^[-*] ", candidate) or re.match(r"^\d+\. ", candidate):
                break
            para_lines.append(candidate.strip())
            i += 1
        text = " ".join(para_lines)
        add_paragraph(doc, text, lead=(text.startswith("本方案将目标定义为：") or text.startswith("将项目第一版定义为")))

    for p in doc.paragraphs:
        p.paragraph_format.widow_control = True
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
